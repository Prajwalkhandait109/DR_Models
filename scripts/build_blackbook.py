"""
Build the completed DR_Blackbook.docx for the project
"Early Detection of Diabetic Retinopathy Using Deep Learning".

All quantitative numbers, training hyperparameters, and per-class metrics in
this report are drawn directly from the notebooks under notebooks/.

Run:
    python scripts/build_blackbook.py
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "blackbook" / "DR_Blackbook.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_centered_heading(doc, text, *, size=16, bold=True, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    return p


def add_para(doc, text, *, bold=False, italic=False, align=None, size=12, space_after=6):
    p = doc.add_paragraph()
    if align is None:
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    return p


def add_chapter(doc, number, title):
    doc.add_page_break()
    # Per template: "Chapter-N" line is TNR 14pt bold, centered
    add_centered_heading(doc, f"Chapter-{number}", size=14, space_after=0)
    # Chapter title is TNR 16pt bold, centered
    add_centered_heading(doc, title, size=16, space_after=12)


def add_section(doc, title, *, level=2):
    # level 2 = main section (e.g. "1.2 Motivation:") -> 14pt bold justify
    # level 3 = subsection -> 13pt bold justify
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(sizes.get(level, 12))


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        if not p.runs:
            r = p.add_run(it)
        else:
            r = p.runs[0]
            r.text = it
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def add_table(doc, headers, rows, *, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for c, w in zip(row.cells, col_widths):
                c.width = w
    doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------
doc = Document()
set_default_font(doc)

# Page setup: Letter size, 1-inch (2.54 cm) margins on all sides per college template
for section in doc.sections:
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


# ===========================================================================
# Title Page
# ===========================================================================
add_centered_heading(doc, "A", size=16, space_after=4)
add_centered_heading(doc, "Project Report on", size=14, bold=False, space_after=18)
add_centered_heading(
    doc,
    "EARLY DETECTION OF DIABETIC RETINOPATHY USING DEEP LEARNING",
    size=18,
    space_after=18,
)
add_para(
    doc,
    "in partial fulfillment for the award of the degree of",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=6,
)
add_centered_heading(doc, "BACHELOR OF TECHNOLOGY", size=16, space_after=4)
add_centered_heading(doc, "IN", size=14, space_after=4)
add_centered_heading(doc, "COMPUTER SCIENCE AND ENGINEERING", size=14, space_after=2)
add_centered_heading(
    doc,
    "(Artificial Intelligence & Machine Learning)",
    size=12,
    bold=False,
    space_after=18,
)
add_para(doc, "Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
for name in [
    "PRAJWAL KHANDAIT",
    "MAHAMADTOHID NAIKWADI",
    "SAMIR MULLA",
    "ADITYA SUTAR",
]:
    add_para(doc, name, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=2)
doc.add_paragraph()
add_para(doc, "Under the Guidance Of", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(
    doc,
    "Mr. Rajesh Kumar",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    bold=True,
    space_after=18,
)
add_centered_heading(doc, "KOLHAPUR INSTITUTE OF TECHNOLOGY'S", size=14, space_after=2)
add_centered_heading(
    doc,
    "COLLEGE OF ENGINEERING (AUTONOMOUS), KOLHAPUR",
    size=14,
    space_after=8,
)
add_para(doc, "2025-26", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)


# ===========================================================================
# Approval Page
# ===========================================================================
doc.add_page_break()
add_centered_heading(doc, "The Project Phase-II Entitled", size=14, bold=False, space_after=14)
add_centered_heading(doc, "On", size=12, bold=False, space_after=10)
add_centered_heading(
    doc,
    "EARLY DETECTION OF DIABETIC RETINOPATHY USING DEEP LEARNING",
    size=16,
    space_after=18,
)
add_para(doc, "Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
for name in [
    "PRAJWAL KHANDAIT",
    "MAHAMADTOHID NAIKWADI",
    "SAMIR MULLA",
    "ADITYA SUTAR",
]:
    add_para(doc, name, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=2)
doc.add_paragraph()
add_para(
    doc,
    "Is approved for the award of the degree of Bachelor of Technology in",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=4,
)
add_centered_heading(doc, "COMPUTER SCIENCE & ENGINEERING (AIML)", size=14, space_after=6)
add_para(doc, "Shivaji University, Kolhapur", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para(doc, "EXAMINERS:", bold=True, space_after=12)
add_para(doc, "Name                                                              Signature")
add_para(doc, "Internal:  __________________________________________________")
add_para(doc, "External: __________________________________________________")
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, "Dr. Uma P. Gurav                                              Dr. Mohan B. Vanarotti", bold=True)
add_para(doc, "HEAD OF THE DEPARTMENT                                   DIRECTOR / PRINCIPAL")


# ===========================================================================
# Certificate
# ===========================================================================
doc.add_page_break()
add_centered_heading(doc, "KOLHAPUR INSTITUTE OF TECHNOLOGY'S", size=16, space_after=2)
add_centered_heading(
    doc,
    "COLLEGE OF ENGINEERING (AUTONOMOUS), KOLHAPUR",
    size=16,
    space_after=18,
)
add_centered_heading(doc, "CERTIFICATE", size=18, space_after=14)
add_para(
    doc,
    "This is to certify that the Project report entitled, "
    "EARLY DETECTION OF DIABETIC RETINOPATHY USING DEEP LEARNING submitted by "
    "Prajwal Khandait (Roll No. A17), Mahamadtohid Naikwadi (Roll No. A78), "
    "Samir Mulla (Roll No. A79), Aditya Sutar (Roll No. A77) in partial "
    "fulfillment for the award of the degree of Bachelor of Technology in "
    "Computer Science and Engineering (Artificial Intelligence and Machine "
    "Learning) at KIT's College of Engineering, Kolhapur, Maharashtra, INDIA, "
    "is a record of their own work carried out under our supervision and guidance.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after=24,
)
doc.add_paragraph()
add_para(doc, "SIGNATURE                                                                      SIGNATURE", space_after=14)
add_para(doc, "DR. UMA P. GURAV                                                       MR. RAJESH KUMAR", bold=True)
add_para(doc, "HEAD OF THE DEPARTMENT                                             SUPERVISOR")
add_para(doc, "Associate Professor                                                          Associate Professor")
add_para(doc, "Department of CSE (AIML & DS)                                       Department of CSE (AIML & DS)")
add_para(doc, "KIT's College of Engineering, Kolhapur                              KIT's College of Engineering, Kolhapur")


# ===========================================================================
# Declaration
# ===========================================================================
doc.add_page_break()
add_centered_heading(doc, "KOLHAPUR INSTITUTE OF TECHNOLOGY'S", size=16, space_after=2)
add_centered_heading(doc, "COLLEGE OF ENGINEERING (AUTONOMOUS), KOLHAPUR", size=16, space_after=18)
add_centered_heading(doc, "DECLARATION", size=18, space_after=14)
add_para(
    doc,
    "We hereby declare that the Project entitled, "
    "EARLY DETECTION OF DIABETIC RETINOPATHY USING DEEP LEARNING submitted to "
    "KIT's College of Engineering, Kolhapur, Maharashtra, INDIA in partial "
    "fulfillment of the award of the Degree of Bachelor of Technology in Computer "
    "Science and Engineering (Artificial Intelligence and Machine Learning) is a "
    "bonafide work carried out by us. The material contained in this Project has "
    "not been submitted to any University or Institution for the award of any degree.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after=24,
)
for name, roll in [
    ("Prajwal Khandait", "A17"),
    ("Mahamadtohid Naikwadi", "A78"),
    ("Samir Mulla", "A79"),
    ("Aditya Sutar", "A77"),
]:
    add_para(doc, f"{name}     (Roll No. {roll})", space_after=2)
doc.add_paragraph()
add_para(doc, "Place: Kolhapur")
add_para(doc, "Date:")


# ===========================================================================
# Acknowledgement
# ===========================================================================
doc.add_page_break()
add_centered_heading(doc, "ACKNOWLEDGEMENT", size=18, space_after=14)
add_para(
    doc,
    "We are deeply grateful to Mr. Rajesh Kumar, Associate Professor, "
    "Department of CSE (AIML), KIT's College of Engineering, Kolhapur, for "
    "their invaluable guidance, constant encouragement, and insightful "
    "suggestions throughout the completion of this project. Their expertise "
    "and support were crucial at every stage of the work.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(
    doc,
    "We also express our sincere thanks to Dr. Uma P. Gurav, Head of the "
    "Department of CSE (AIML & DS), for providing the necessary facilities and "
    "a conducive environment to carry out this project.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(
    doc,
    "Our heartfelt thanks to Dr. Mohan B. Vanarotti, Director / Principal of "
    "KIT's College of Engineering, Kolhapur, for their constant motivation and support.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(
    doc,
    "We would also like to thank all the faculty members and staff of the "
    "Department of CSE (AIML & DS) for their cooperation and assistance. "
    "Finally, we are thankful to our friends and family for their encouragement, "
    "patience, and continuous support during the course of this project.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after=18,
)
for name, roll in [
    ("Prajwal Khandait", "A17"),
    ("Mahamadtohid Naikwadi", "A78"),
    ("Samir Mulla", "A79"),
    ("Aditya Sutar", "A77"),
]:
    add_para(doc, f"{name}     (Roll No. {roll})", space_after=2)
doc.add_paragraph()
add_para(doc, "Place: Kolhapur")
add_para(doc, "Date:")


# ===========================================================================
# Declaration of Plagiarism
# ===========================================================================
doc.add_page_break()
add_centered_heading(doc, "KOLHAPUR INSTITUTE OF TECHNOLOGY'S", size=16, space_after=2)
add_centered_heading(doc, "COLLEGE OF ENGINEERING (AUTONOMOUS), KOLHAPUR", size=16, space_after=18)
add_centered_heading(doc, "Declaration of Plagiarism", size=16, space_after=14)
add_para(
    doc,
    "We hereby declare that the dissertation entitled "
    "\"EARLY DETECTION OF DIABETIC RETINOPATHY USING DEEP LEARNING\", which is "
    "being submitted to KIT's College of Engineering, Kolhapur, Maharashtra, "
    "INDIA, in partial fulfillment of the B.Tech. CSE (AIML) course, is a "
    "complete and exclusive report of our work, conducted honestly. We further "
    "declare that the content and the material has not been used by anyone in "
    "the past, and has not been submitted to any educational institution or any "
    "University for any degree award program. The material content of this "
    "thesis is duly acknowledged. Plagiarism similarity index check has been "
    "performed and the report has been attached separately.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after=24,
)
add_para(doc, "Date:  /   / 2026                                                                     Signature of Students:")
add_para(doc, "Place: Kolhapur                                                                       Names: Prajwal Khandait,")
add_para(doc, "                                                                                                       Mahamadtohid Naikwadi,")
add_para(doc, "                                                                                                       Samir Mulla, Aditya Sutar")
doc.add_paragraph()
add_para(doc, "Date:  /   / 2026                                                                     Signature of the Guide:")
add_para(doc, "Place: Kolhapur                                                                       Name: Mr. Rajesh Kumar")


# ===========================================================================
# Index / Table of Contents
# ===========================================================================
doc.add_page_break()
add_centered_heading(doc, "INDEX", size=18, space_after=14)

toc_rows = [
    ("Abstract", "i"),
    ("List of Figures", "ii"),
    ("List of Tables", "iii"),
    ("List of Abbreviations", "iv"),
    ("Chapter 1: INTRODUCTION", "1"),
    ("    1.1 Introduction", "2"),
    ("    1.2 Motivation", "3"),
    ("    1.3 Report Organization", "4"),
    ("Chapter 2: PROBLEM STATEMENT", "5"),
    ("    2.1 Literature Survey", "6"),
    ("    2.2 Need of Work", "10"),
    ("    2.3 Problem Statement", "11"),
    ("    2.4 Objectives", "12"),
    ("Chapter 3: DESIGN DETAILS", "13"),
    ("    3.1 System Architecture", "14"),
    ("    3.2 Design Methodology", "16"),
    ("    3.3 System Design Diagrams", "17"),
    ("    3.4 Dataset Design", "20"),
    ("Chapter 4: IMPLEMENTATION DETAILS", "22"),
    ("    4.1 Development Environment", "23"),
    ("    4.2 Module-wise Implementation", "24"),
    ("    4.3 Model 1: InceptionResNetV2", "26"),
    ("    4.4 Model 2: Swin Transformer", "29"),
    ("    4.5 Model 3: Vision Transformer", "32"),
    ("    4.6 Model 4: Two-Stage Screening Pipeline", "35"),
    ("    4.7 System Workflow", "38"),
    ("Chapter 5: RESULTS AND DISCUSSION", "40"),
    ("    5.1 Dataset and Experimental Setup", "41"),
    ("    5.2 Evaluation Metrics", "42"),
    ("    5.3 InceptionResNetV2 Results", "43"),
    ("    5.4 Swin Transformer Results", "45"),
    ("    5.5 Vision Transformer Results", "47"),
    ("    5.6 Two-Stage Screening Results", "49"),
    ("    5.7 Comparative Analysis", "51"),
    ("Chapter 6: CONCLUSION AND FUTURE SCOPE", "53"),
    ("Chapter 7: PAPERS PUBLISHED", "55"),
    ("Chapter 8: REFERENCES", "56"),
]
add_table(doc, ["Topic", "Page No."], toc_rows, col_widths=[Cm(12), Cm(3)])


# ===========================================================================
# Lists of Figures / Tables / Abbreviations
# ===========================================================================
doc.add_page_break()
add_centered_heading(doc, "List of Figures", size=16, space_after=10)
fig_rows = [
    ("1", "System Architecture Diagram"),
    ("2", "Data Flow Diagram (DFD)"),
    ("3", "Activity Diagram"),
    ("4", "Use Case Diagram"),
    ("5", "Sequence Diagram"),
    ("6", "APTOS 2019 Dataset Class Distribution"),
    ("7", "Sample Fundus Images (one per class)"),
    ("8", "InceptionResNetV2 — Training & Validation Accuracy / Loss"),
    ("9", "InceptionResNetV2 — Confusion Matrix"),
    ("10", "InceptionResNetV2 — ROC-AUC Curves"),
    ("11", "InceptionResNetV2 — GradCAM Heatmap"),
    ("12", "Swin Transformer — Training Loss & Validation Accuracy"),
    ("13", "Swin Transformer — Confusion Matrix"),
    ("14", "Swin Transformer — ROC-AUC Curves"),
    ("15", "Swin Transformer — GradCAM Samples (Correct / Wrong)"),
    ("16", "Vision Transformer — Training Curves"),
    ("17", "Vision Transformer — Confusion Matrix"),
    ("18", "Vision Transformer — ROC-AUC Curves"),
    ("19", "Vision Transformer — GradCAM++ Samples"),
    ("20", "Two-Stage Pipeline — Block Diagram"),
    ("21", "Two-Stage Pipeline — Stage 2 Confusion Matrix"),
    ("22", "Two-Stage Pipeline — Stage 2 ROC-AUC Curves"),
    ("23", "Comparative Bar Chart of Model Accuracies"),
]
add_table(doc, ["Sr. No.", "Name of Figure"], fig_rows, col_widths=[Cm(2), Cm(13)])

add_centered_heading(doc, "List of Tables", size=16, space_after=10)
tab_rows = [
    ("2.1", "Literature Review Summary"),
    ("3.1", "APTOS 2019 Dataset Class Distribution"),
    ("4.1", "Hyperparameters per Model"),
    ("5.1", "Training Parameters per Model"),
    ("5.2", "Testing Parameters per Model"),
    ("5.3", "InceptionResNetV2 — Per-Class Metrics"),
    ("5.4", "InceptionResNetV2 — Overall Performance"),
    ("5.5", "Swin Transformer — Per-Class Metrics"),
    ("5.6", "Swin Transformer — Overall Performance"),
    ("5.7", "Vision Transformer — Per-Class Metrics"),
    ("5.8", "Vision Transformer — Overall Performance"),
    ("5.9", "Two-Stage Pipeline — Stage 2 Per-Class Metrics"),
    ("5.10", "Two-Stage Pipeline — Overall Performance"),
    ("5.11", "Comparative Performance of All Models"),
]
add_table(doc, ["Sr. No.", "Name of Table"], tab_rows, col_widths=[Cm(2), Cm(13)])

add_centered_heading(doc, "List of Abbreviations", size=16, space_after=10)
abbr_rows = [
    ("01", "DR", "Diabetic Retinopathy"),
    ("02", "CNN", "Convolutional Neural Network"),
    ("03", "ViT", "Vision Transformer"),
    ("04", "Swin", "Shifted Window Transformer"),
    ("05", "DFD", "Data Flow Diagram"),
    ("06", "CLAHE", "Contrast Limited Adaptive Histogram Equalization"),
    ("07", "AUC-ROC", "Area Under Curve — Receiver Operating Characteristic"),
    ("08", "NPDR", "Non-Proliferative Diabetic Retinopathy"),
    ("09", "PDR", "Proliferative Diabetic Retinopathy"),
    ("10", "mAP", "Mean Average Precision"),
    ("11", "SGD", "Stochastic Gradient Descent"),
    ("12", "API", "Application Programming Interface"),
    ("13", "GPU", "Graphics Processing Unit"),
    ("14", "IRV2", "Inception ResNet Version 2"),
    ("15", "APTOS", "Asia Pacific Tele-Ophthalmology Society"),
    ("16", "GradCAM", "Gradient-weighted Class Activation Mapping"),
    ("17", "MLP", "Multi-Layer Perceptron"),
    ("18", "AdamW", "Adam Optimizer with Decoupled Weight Decay"),
    ("19", "EHR", "Electronic Health Record"),
    ("20", "OvR", "One-vs-Rest"),
]
add_table(doc, ["Sr. No.", "Abbreviation", "Full Form"], abbr_rows, col_widths=[Cm(2), Cm(3.5), Cm(10)])


# ===========================================================================
# Abstract
# ===========================================================================
doc.add_page_break()
add_centered_heading(doc, "ABSTRACT", size=18, space_after=14)
add_para(
    doc,
    "Diabetic Retinopathy (DR) is one of the leading causes of preventable "
    "blindness among working-age adults globally. Conventional DR screening "
    "depends on manual examination of retinal fundus images by trained "
    "ophthalmologists, which is time-consuming, subjective, and inaccessible "
    "in resource-limited regions. This project presents a deep-learning-based "
    "system that automatically classifies retinal fundus images from the "
    "APTOS 2019 Blindness Detection dataset (3,662 images) into five clinically "
    "recognized severity classes: No DR, Mild NPDR, Moderate NPDR, Severe NPDR, "
    "and Proliferative DR.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(
    doc,
    "Four deep architectures were independently implemented, trained, and "
    "evaluated: (i) InceptionResNetV2 with multistage transfer learning; "
    "(ii) Swin Transformer (swin_base_patch4_window7_224) with hierarchical "
    "shifted-window self-attention; (iii) Vision Transformer "
    "(vit_base_patch16_224) with single-phase fine-tuning, weighted random "
    "sampling, and label smoothing; and (iv) a two-stage cascaded pipeline "
    "combining an EfficientNetV2B0 binary screener with a Swin Transformer "
    "severity grader. Each model was evaluated on an 80/20 stratified validation "
    "split using accuracy, precision, recall, F1-score, ROC-AUC, and confusion "
    "matrices. GradCAM / GradCAM++ visualizations were generated to provide "
    "lesion-level interpretability of the predictions.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(
    doc,
    "Among the four implementations, the Vision Transformer achieved the "
    "highest overall validation accuracy of 85.40 % with a weighted F1 of "
    "0.85, followed by the Swin Transformer (83 %), the Two-Stage cascade "
    "(70 % on the 4-class severity sub-problem), and InceptionResNetV2 "
    "(79 %). A Flask-based web interface allows clinicians to upload retinal "
    "images and obtain class predictions with GradCAM explainability — "
    "supporting clinically-aligned, transparent, and scalable DR screening.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)


# ===========================================================================
# CHAPTER 1 — INTRODUCTION
# ===========================================================================
add_chapter(doc, 1, "INTRODUCTION")
add_section(doc, "1.1 Introduction", level=2)
add_para(
    doc,
    "Diabetic Retinopathy (DR) is a severe microvascular complication of diabetes "
    "mellitus that damages the blood vessels in the light-sensitive tissue at the "
    "back of the eye — the retina. It is the leading cause of new cases of "
    "blindness and vision impairment among working-age adults globally, "
    "affecting approximately 80–85 % of patients who have had diabetes for more "
    "than 10 years. According to the International Diabetes Federation, diabetes "
    "affects over 463 million people worldwide, a number projected to reach 700 "
    "million by 2045, making the scale of DR a growing public health crisis.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(
    doc,
    "Traditional DR diagnosis relies on the manual examination of retinal fundus "
    "images by trained ophthalmologists — a process that is time-consuming, "
    "expensive, subjective, and dependent on expert availability. Early detection "
    "is critical because timely treatment can prevent up to 95 % of vision loss "
    "from DR. Consequently, there is a pressing need for automated, scalable, and "
    "clinically reliable DR screening solutions.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(
    doc,
    "The proposed Diabetic Retinopathy Detection System using Deep Learning "
    "automates the identification and classification of DR from retinal fundus "
    "images using advanced Convolutional Neural Network and Transformer-based "
    "techniques. Four complementary architectures — InceptionResNetV2, Swin "
    "Transformer, Vision Transformer, and a two-stage EfficientNetV2B0 + Swin "
    "cascade — are independently trained and benchmarked on the APTOS 2019 "
    "dataset. The best performing model is exposed through a Flask web "
    "application that delivers predictions and GradCAM explainability heatmaps "
    "in real time.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "1.2 Motivation", level=2)
add_para(
    doc,
    "The rapid increase in global diabetes prevalence, combined with a severe "
    "shortage of trained ophthalmologists, has made manual DR screening both "
    "inadequate and unsustainable. Deep learning models have demonstrated "
    "performance comparable to expert ophthalmologists in several landmark "
    "studies, achieving AUC scores above 0.99 on large retinal image datasets. "
    "This technological capability, combined with the urgent clinical need, "
    "makes automated DR detection both feasible and essential.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc, "Current Screening Challenges:", bold=True)
add_bullets(doc, [
    "Manual Examination — DR diagnosis traditionally relies on manual fundus image examination, which is labor-intensive, time-consuming, and prone to inter-observer variability.",
    "Specialist Shortage — A critical global shortage of ophthalmologists means demand for diabetic eye screening far exceeds available expertise, especially in low- and middle-income countries.",
    "Accessibility Gap — Rural and underserved communities lack access to specialized eye care facilities, resulting in delayed diagnoses and preventable vision loss.",
    "Scaling Healthcare — With increasing diabetes prevalence, healthcare systems cannot scale manual screening processes quickly enough to meet the growing demand.",
])

add_section(doc, "1.3 Report Organization", level=2)
add_bullets(doc, [
    "Chapter 1 — Introduction: motivation, problem context, and report structure.",
    "Chapter 2 — Problem Statement: literature survey, identified gaps, problem definition, and objectives.",
    "Chapter 3 — Design Details: system architecture, design methodology, UML diagrams, and dataset design.",
    "Chapter 4 — Implementation Details: development environment, module-wise implementation, and per-model implementation (InceptionResNetV2, Swin, ViT, Two-Stage).",
    "Chapter 5 — Results and Discussion: per-model and comparative evaluation including confusion matrices, ROC-AUC, classification reports, and GradCAM analyses.",
    "Chapter 6 — Conclusion and Future Scope: summary of findings and directions for future enhancement.",
    "Chapter 7 — Papers Published / Progress Reports.",
    "Chapter 8 — References.",
])


# ===========================================================================
# CHAPTER 2 — PROBLEM STATEMENT
# ===========================================================================
add_chapter(doc, 2, "PROBLEM STATEMENT")
add_para(
    doc,
    "This chapter presents a detailed review of recent research works related to "
    "automated Diabetic Retinopathy detection and grading using deep learning "
    "techniques. Special emphasis is placed on Convolutional Neural Networks "
    "(CNNs), hybrid CNN-residual networks, and transformer-based architectures, "
    "all of which form the basis of the four models implemented in this project.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "2.1 Literature Survey", level=2)
add_para(
    doc,
    "In recent years, deep learning has significantly advanced automated DR "
    "detection systems, with numerous studies leveraging pretrained CNN models "
    "and, more recently, vision transformers to achieve clinical-grade performance. "
    "The literature most relevant to this project is summarized in Table 2.1.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

lit_rows = [
    ("Guefrachi et al. (2025) [1]",
     "DR Detection Using Deep Learning Multistage Training",
     "InceptionResNetV2, VGG16, VGG19, DenseNet121, MobileNetV2; 2-stage fine-tuning (frozen LR=1e-3 then fine-tune LR=1e-5, 25 epochs each); 97.8% accuracy on APTOS."),
    ("Balasamy & Suganyadevi (2025) [2]",
     "Multi-Dimensional Fuzzy Based DR Detection through Deep CNN",
     "Fuzzy entropy multi-dimensional thresholding + Deep CNN on 42,000 synthetic retinal images; Monte Carlo validation; precise DR severity detection."),
    ("Chandra, Tiwari et al. [3]",
     "DR Prediction Based on CNN and AlexNet",
     "CNN 97% and AlexNet 93% accuracy on APTOS; 50×50 image resolution; balanced 1,737 samples/class."),
    ("Gulshan et al. (2016) [4]",
     "DL Algorithm for DR Detection in Fundus Photos",
     "Inception-v3 on 128,175 retinal images; AUC 0.991; performance comparable to expert ophthalmologists."),
    ("Abramoff et al. (2018) [5]",
     "Autonomous AI Diagnostic System (IDx-DR)",
     "DenseNet-based IDx-DR; 87.2% sensitivity, 90.7% specificity; first FDA-cleared autonomous AI diagnostic for DR."),
    ("Pratt et al. (2016) [6]",
     "CNN for Diabetic Retinopathy",
     "CNN on Kaggle 2015 DR; 75.6% accuracy for 5-class severity grading."),
    ("Chetoui & Akhloufi (2020) [7]",
     "Explainable End-to-End DL for DR Grading",
     "EfficientNet + Grad-CAM visual explainability; competitive accuracy with clinical interpretability."),
    ("Solanki et al. (2023) [8]",
     "ResNet50 with Attention for DR Severity",
     "ResNet50 + attention on MESSIDOR; 94.7% accuracy; highlights microaneurysms, hemorrhages, exudates."),
    ("Li et al. (2019) [9]",
     "PathoEye DL Framework",
     "Multi-label deep learning for DR and co-occurring retinal pathologies."),
    ("Wan et al. (2018) [10]",
     "Deep CNN for DR Detection",
     "Deep CNN on MESSIDOR for binary and multi-class DR classification."),
    ("Dosovitskiy et al. (2021) [11]",
     "An Image is Worth 16x16 Words — ViT",
     "Vision Transformer (ViT) with patch-wise self-attention; outperforms CNNs on large pretraining datasets."),
    ("Liu et al. (2021) [12]",
     "Swin Transformer",
     "Hierarchical shifted-window self-attention; linear complexity, multi-scale features, strong on dense prediction."),
    ("Tan & Le (2021) [13]",
     "EfficientNetV2",
     "Compound-scaled CNN with Fused-MBConv; trains faster with smaller models — ideal for screening stage."),
]
add_para(doc, "Table 2.1: Literature Review Summary", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table(doc, ["Authors", "Title", "Methodology"], lit_rows, col_widths=[Cm(4), Cm(4.5), Cm(7.5)])

add_section(doc, "Research Gaps Identified", level=3)
add_bullets(doc, [
    "RG1 — Limited dataset diversity: most studies use small or single-source datasets and do not capture real-world variability.",
    "RG2 — Binary vs. multi-class classification: many proposed solutions perform only binary DR/No-DR detection rather than 5-class grading.",
    "RG3 — Lack of model explainability: most high-accuracy models function as black boxes without lesion-level visual interpretation.",
    "RG4 — Limited real-world validation: many solutions are validated in controlled conditions and not exposed through deployable interfaces.",
    "RG5 — Single-architecture studies: comparative evaluation of CNN, hybrid, and transformer architectures on the same dataset under identical preprocessing is rare.",
])

add_section(doc, "2.2 Need of Work", level=2)
add_para(
    doc,
    "The increasing global burden of diabetes and the consequent rise in DR cases "
    "have made traditional manual screening systems inadequate and unsustainable. "
    "With an estimated 93 million people worldwide affected by some form of DR "
    "and only one ophthalmologist per 100,000 people in many developing nations, "
    "there is an urgent need for automated, intelligent DR screening systems. "
    "A comparative deep-learning study using a single standardized dataset "
    "(APTOS 2019) and consistent evaluation protocol provides the evidence "
    "needed to choose the most clinically suitable architecture.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "2.3 Problem Statement", level=2)
add_para(
    doc,
    "Diagnosing Diabetic Retinopathy requires expert ophthalmologists and manual "
    "examination of retinal fundus images, which is time-consuming, subjective, "
    "and inaccessible in many regions. This project aims to develop an AI-based "
    "system using deep learning for automated detection and classification of DR "
    "into five clinically recognized severity levels: No DR, Mild NPDR, Moderate "
    "NPDR, Severe NPDR, and Proliferative DR. Four independent architectures are "
    "implemented and compared, and the best-performing model is integrated into "
    "a Flask web interface for clinical deployment.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "2.4 Objectives", level=2)
add_bullets(doc, [
    "Obj1 — Build a reproducible pipeline on the APTOS 2019 dataset with stratified 80/20 split and consistent augmentation.",
    "Obj2 — Implement 5-class DR severity grading rather than binary detection.",
    "Obj3 — Train and benchmark four architectures: InceptionResNetV2, Swin Transformer, Vision Transformer, and a two-stage EfficientNetV2B0 + Swin cascade.",
    "Obj4 — Evaluate all models using accuracy, precision, recall, F1-score, ROC-AUC, and confusion matrices on an identical validation split.",
    "Obj5 — Provide lesion-level explainability via GradCAM / GradCAM++ heatmaps for clinical interpretability.",
    "Obj6 — Deploy the best model through a Flask web interface for real-world image upload, prediction, and result visualization.",
])


# ===========================================================================
# CHAPTER 3 — DESIGN DETAILS
# ===========================================================================
add_chapter(doc, 3, "DESIGN DETAILS")

add_section(doc, "3.1 System Architecture", level=2)
add_para(doc, "Fig 3.1: System Architecture of DR Detection System", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    doc,
    "The system architecture is designed as a modular and scalable pipeline. It "
    "comprises six logical layers: (1) image acquisition through the web "
    "interface; (2) preprocessing (resize, normalization, optional CLAHE); "
    "(3) inference via the trained deep learning model (any of four available "
    "models); (4) post-processing (softmax → class label → clinical "
    "recommendation); (5) explainability (GradCAM / GradCAM++ heatmap "
    "generation); and (6) result visualization on the Flask web interface.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "3.2 Design Methodology", level=2)
add_para(
    doc,
    "The project follows an agile and iterative methodology emphasizing "
    "modularity, feedback-driven improvement, and continuous integration of "
    "experiments across multiple architectures.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "3.2.1 Planning Phase", level=3)
add_para(doc,
    "Project goals were defined with a focus on clinical-grade DR detection. "
    "The APTOS 2019 dataset was selected. Tooling — TensorFlow/Keras, PyTorch, "
    "timm, pytorch-grad-cam, tf-explain, Flask, Google Colab, and Kaggle "
    "Notebooks — was finalized.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "3.2.2 Risk Analysis Phase", level=3)
add_para(doc,
    "Identified risks include class imbalance in the APTOS dataset, hardware "
    "limitations during training, overfitting of high-capacity transformer "
    "models, and Stage-1 error propagation in the two-stage cascade. Mitigation "
    "strategies include data augmentation, weighted random sampling, label "
    "smoothing, dropout regularization, gradient clipping, and multistage "
    "fine-tuning with low learning rates.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "3.2.3 Development Phase", level=3)
add_para(doc,
    "Each model was implemented in an independent Jupyter notebook under "
    "notebooks/. Google Colab T4 GPU and Kaggle Tesla T4 (15.6 GB VRAM) were "
    "used for training. Hyperparameter tuning and validation were performed "
    "iteratively per architecture.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "3.2.4 Evaluation Phase", level=3)
add_para(doc,
    "All four models were evaluated on the same 80/20 stratified validation "
    "split using accuracy, precision, recall, F1-score, AUC-ROC, confusion "
    "matrices, and GradCAM / GradCAM++ visualizations.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "3.3 System Design Diagrams", level=2)
add_section(doc, "3.3.1 Data Flow Diagram (DFD)", level=3)
add_para(doc, "Fig 3.2: Data Flow Diagram", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,
    "The DFD illustrates the workflow of the DR detection system. The user "
    "uploads a retinal fundus image via the web interface. The image is "
    "preprocessed (resize, normalization), then fed into the selected model "
    "for inference. The model outputs the predicted DR class with a confidence "
    "score, which is displayed alongside a GradCAM heatmap and a clinical "
    "recommendation.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "3.3.2 Activity Diagram", level=3)
add_para(doc, "Fig 3.3: Activity Diagram", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,
    "The activity diagram outlines the full DR detection process: image upload "
    "→ preprocessing → model inference → DR detected? → display severity and "
    "recommendation. If image quality is insufficient, the user is prompted to "
    "upload a clearer image.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "3.3.3 Use Case Diagram", level=3)
add_para(doc, "Fig 3.4: Use Case Diagram", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,
    "The use case diagram shows how Patient, Doctor, and Admin actors interact "
    "with the system. Key use cases include Upload Fundus Image, View DR "
    "Classification Result, View GradCAM Heatmap, Download Report, and Manage "
    "Users / Monitor System Performance.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "3.3.4 Sequence Diagram", level=3)
add_para(doc, "Fig 3.5: Sequence Diagram", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,
    "The sequence diagram illustrates the interaction flow between the User, "
    "Browser, Flask Backend, Model Inference Module, and Storage. The user "
    "uploads an image; Flask receives the POST request, invokes the chosen "
    "model, computes the GradCAM heatmap, persists results, and returns the "
    "classification and visualization to the browser.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "3.4 Dataset Design", level=2)
add_section(doc, "3.4.1 Dataset Description", level=3)
add_para(doc, "Fig 3.6: APTOS 2019 Class Distribution", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,
    "The APTOS 2019 dataset (Kaggle: sovitrath/diabetic-retinopathy-224x224-2019-data) "
    "contains 3,662 retinal fundus images graded by certified ophthalmologists "
    "into five DR severity classes. The dataset reflects real-world class "
    "imbalance, with No DR being the dominant class. Images were captured "
    "under varying lighting conditions and camera settings, reflecting "
    "clinical diversity. A stratified 80/20 split was used for training and "
    "validation (2,929 training, 733 validation).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc, "Table 3.1: APTOS 2019 Class Distribution", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
ds_rows = [
    ("0", "No DR", "1,805", "49.3 %"),
    ("1", "Mild NPDR", "370", "10.1 %"),
    ("2", "Moderate NPDR", "999", "27.3 %"),
    ("3", "Severe NPDR", "193", "5.3 %"),
    ("4", "Proliferative DR", "295", "8.0 %"),
    ("—", "Total", "3,662", "100.0 %"),
]
add_table(doc, ["Class ID", "Severity Class", "Image Count", "Proportion"], ds_rows, col_widths=[Cm(2), Cm(5), Cm(3), Cm(3)])

add_section(doc, "3.4.2 Sample Dataset View", level=3)
add_para(doc, "Fig 3.7: Sample Fundus Images (one per class)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,
    "Healthy retinas (No DR) appear with clear, uniform vasculature. Mild DR "
    "shows occasional microaneurysms. Moderate DR presents hemorrhages and "
    "exudates. Severe DR exhibits extensive hemorrhages across all quadrants. "
    "Proliferative DR shows neovascularization — new abnormal blood vessels "
    "growing from the retinal surface.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)


# ===========================================================================
# CHAPTER 4 — IMPLEMENTATION DETAILS
# ===========================================================================
add_chapter(doc, 4, "IMPLEMENTATION DETAILS")

add_section(doc, "4.1 Development Environment", level=2)
add_section(doc, "4.1.1 Hardware Requirements", level=3)
add_bullets(doc, [
    "Platform: Google Colab and Kaggle Notebooks (cloud-based)",
    "GPU: NVIDIA Tesla T4 (15.6 GB VRAM) on both platforms",
    "RAM: 12–16 GB",
    "Storage: ~78 GB (Colab) and ~73 GB (Kaggle)",
])
add_section(doc, "4.1.2 Software Requirements", level=3)
add_bullets(doc, [
    "Programming Language: Python 3.10 / 3.11",
    "Operating System: Windows 10/11 (local), Linux (Colab/Kaggle)",
    "IDEs: Jupyter Notebook, Google Colab, VS Code",
    "Web Stack: Flask 2.3+, HTML, CSS, JavaScript",
    "Deep Learning Frameworks: TensorFlow 2.14+ / Keras (InceptionResNetV2, EfficientNetV2B0), PyTorch 2.10 + timm (Swin, ViT)",
])
add_section(doc, "4.1.3 Libraries / Frameworks Used", level=3)
add_bullets(doc, [
    "tensorflow / keras — InceptionResNetV2 and EfficientNetV2B0 training",
    "torch, torchvision, timm — Swin Transformer and Vision Transformer training",
    "pytorch-grad-cam — GradCAM and GradCAM++ heatmap generation",
    "tf-explain — GradCAM for the InceptionResNetV2 model",
    "opencv-python, Pillow — Image I/O and preprocessing",
    "pandas, numpy — Data manipulation and numerical computations",
    "scikit-learn — Train/validation split, confusion matrix, classification report, ROC-AUC",
    "matplotlib, seaborn — Visualization of metrics, confusion matrix heatmaps, ROC curves",
    "flask — Web application backend exposing the inference API",
    "kagglehub — Programmatic dataset download from Kaggle",
])

add_section(doc, "4.2 Module-wise Implementation", level=2)
add_section(doc, "4.2.1 Data Collection & Preprocessing Module", level=3)
add_para(doc,
    "Retinal fundus images and corresponding severity labels are loaded from "
    "the APTOS 2019 Kaggle dataset using the kagglehub API. A pandas DataFrame "
    "is constructed mapping each image filepath to its class. A stratified "
    "80/20 train/validation split is generated using scikit-learn's "
    "train_test_split. Per-framework preprocessing pipelines apply resize "
    "to 224×224 plus framework-specific normalization (InceptionResNetV2's "
    "preprocess_input [-1,1] for TF models, ImageNet mean/std for PyTorch "
    "models). Augmentations vary per model and are detailed in §4.3–§4.6.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "4.2.2 Model Training Modules", level=3)
add_para(doc,
    "Four distinct training modules were implemented, one per architecture. "
    "Each module trains, validates, checkpoints, and saves the final model. "
    "Hyperparameters per model are summarized in Table 4.1.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

hp_rows = [
    ("InceptionResNetV2", "TF / Keras", "Adam → Adam", "1e-3 → 1e-5", "5 + 15", "32", "Categorical CE", "5"),
    ("Swin Transformer", "PyTorch / timm", "Adam → Adam", "1e-3 → 1e-5", "5 + 15", "32", "Cross-Entropy", "5"),
    ("Vision Transformer", "PyTorch / timm", "AdamW + Cosine", "3e-5 (warmup 3 ep)", "15", "32", "CE + Label Smoothing 0.1", "5"),
    ("Two-Stage (EffNetV2B0 + Swin)", "TF + PyTorch", "Adam / Adam → Adam", "1e-3 → 1e-5", "5+10 (S1), 5+15 (S2)", "32", "BCE (S1), CE (S2)", "2 (S1) + 4 (S2)"),
]
add_para(doc, "Table 4.1: Hyperparameters per Model", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table(doc, ["Model", "Framework", "Optimizer", "Learning Rate", "Epochs", "Batch Size", "Loss", "Classes"], hp_rows)

add_section(doc, "4.2.3 Inference Module", level=3)
add_para(doc,
    "The Flask backend loads the trained models on startup. The /predict "
    "endpoint accepts POST requests with an uploaded fundus image, runs the "
    "selected model, computes the softmax probabilities, generates the "
    "corresponding GradCAM / GradCAM++ heatmap, and returns the predicted "
    "class, confidence, and heatmap URL as JSON.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "4.2.4 Web Interface Module", level=3)
add_para(doc,
    "A lightweight Flask web application provides image upload, prediction "
    "results, color-coded severity badges (green for No DR, yellow for "
    "Mild/Moderate, red for Severe/Proliferative), confidence percentage, "
    "GradCAM overlay, and a clinical recommendation. The interface uses "
    "asynchronous JavaScript fetch calls to avoid page reload.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

# 4.3 InceptionResNetV2
add_section(doc, "4.3 Model 1 — InceptionResNetV2 (Multistage Fine-Tuning)", level=2)
add_section(doc, "4.3.1 Architecture", level=3)
add_para(doc,
    "InceptionResNetV2 fuses the multi-scale feature extraction of Inception "
    "modules (parallel 1×1, 3×3, 5×5 paths) with the residual skip connections "
    "of ResNet, enabling stable training of a ~164-layer network. The "
    "Conv2d_7b_ac layer produces 1,536-dimensional spatial features at 8×8 "
    "resolution, which feed the classification head.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "4.3.2 Classification Head", level=3)
add_bullets(doc, [
    "GlobalAveragePooling2D → 1,536-d feature vector",
    "Dropout(0.5)",
    "Dense(5, activation='softmax')",
])
add_section(doc, "4.3.3 Preprocessing & Augmentation", level=3)
add_bullets(doc, [
    "preprocess_input — InceptionResNetV2-specific normalization to [-1, 1]",
    "Horizontal flip",
    "Random rotation ±15°",
    "Random zoom 0.2",
])
add_section(doc, "4.3.4 Training Strategy", level=3)
add_bullets(doc, [
    "Stage 1: classification head only (backbone frozen) — Adam, LR = 1e-3, 5 epochs",
    "Stage 2: full fine-tuning (all layers trainable) — Adam, LR = 1e-5, 15 epochs",
    "Callbacks: ModelCheckpoint(monitor='val_accuracy'), EarlyStopping(patience=5, monitor='val_loss', restore_best_weights=True)",
    "Loss: categorical_crossentropy",
    "Saved artifacts: best_inceptionresnetv2_checkpoint.h5, inceptionresnetv2_dr_model.h5",
])
add_section(doc, "4.3.5 GradCAM Target Layer", level=3)
add_bullets(doc, ["Target layer: conv_7b_ac (final activated convolutional layer, 8×8 spatial map)"])

# 4.4 Swin Transformer
add_section(doc, "4.4 Model 2 — Swin Transformer (swin_base_patch4_window7_224)", level=2)
add_section(doc, "4.4.1 Architecture", level=3)
add_para(doc,
    "Swin Transformer is a hierarchical vision transformer that performs "
    "self-attention within local non-overlapping windows. By alternating "
    "regular and shifted windows across consecutive blocks, the model achieves "
    "cross-window information flow while maintaining linear complexity with "
    "respect to image size. Key properties: 4 hierarchical stages with patch "
    "merging, embed dim 128, heads [4, 8, 16, 32], depths [2, 2, 18, 2], "
    "MLP ratio 4.0. The original classification head is replaced with a 5-class "
    "linear layer.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "4.4.2 Preprocessing & Augmentation", level=3)
add_bullets(doc, [
    "Resize(224, 224)",
    "RandomHorizontalFlip",
    "RandomRotation(±10°)",
    "Normalize (mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])",
])
add_section(doc, "4.4.3 Training Strategy", level=3)
add_bullets(doc, [
    "Stage 1: head-only training — Adam, LR = 1e-3, 5 epochs",
    "Stage 2: full fine-tuning — Adam, LR = 1e-5, 15 epochs (total 20 epochs)",
    "Loss: CrossEntropyLoss",
    "Device: CUDA (Tesla T4) when available",
    "Saved artifacts: swin_checkpoint.pth, swin_transformer_dr.pth",
])
add_section(doc, "4.4.4 GradCAM Target Layer", level=3)
add_bullets(doc, ["Target layer: model.layers[-1].blocks[-1].norm1"])

# 4.5 Vision Transformer
add_section(doc, "4.5 Model 3 — Vision Transformer (vit_base_patch16_224)", level=2)
add_section(doc, "4.5.1 Architecture", level=3)
add_para(doc,
    "ViT-Base/16 applies multi-head self-attention directly over 16×16 image "
    "patches. The 224×224 input is split into 196 patches plus 1 CLS token "
    "(197 tokens), each projected to a 768-dimensional embedding. Twelve "
    "transformer encoder blocks (12 attention heads each, MLP ratio 4.0) "
    "process all tokens with global attention from layer 1.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc, "Custom Classification Head:", bold=True)
add_bullets(doc, [
    "LayerNorm(768)",
    "Dropout(0.30)",
    "Linear(768 → 512) + GELU",
    "Dropout(0.15)",
    "Linear(512 → 5)",
])
add_para(doc, "Total parameters: 86,196,485")
add_section(doc, "4.5.2 Preprocessing & Augmentation", level=3)
add_bullets(doc, [
    "Resize(256) → RandomCrop(224)",
    "RandomHorizontalFlip (p=0.5)",
    "RandomVerticalFlip (p=0.3)",
    "ColorJitter(brightness=0.2, contrast=0.2)",
    "RandomRotation(±15°)",
    "RandomAffine(shear=10°)",
    "Normalize (ImageNet mean/std)",
    "WeightedRandomSampler — inverse-frequency class weights for balanced batches",
])
add_section(doc, "4.5.3 Training Strategy", level=3)
add_bullets(doc, [
    "Optimizer: AdamW (decoupled weight decay)",
    "Learning Rate: 3e-5 (warmup 3 epochs from 3e-6) → CosineAnnealingLR (T_max=12, eta_min=1e-6)",
    "Weight Decay: 1e-2",
    "Loss: CrossEntropyLoss with label_smoothing=0.1",
    "Gradient clipping: max_norm=1.0",
    "Early stopping patience: 10",
    "Max epochs: 15 (plus 2 additional fine-tune epochs from best checkpoint)",
    "Best validation accuracy: 85.40 % (epoch 14 of base + 2 of phase 2)",
])
add_section(doc, "4.5.4 GradCAM++ Target Layer & Reshape", level=3)
add_para(doc,
    "GradCAM++ is applied at model.backbone.blocks[-1].norm1. A reshape "
    "transform removes the CLS token and reshapes the remaining 196 patch "
    "tokens into a 14×14 spatial activation grid for heatmap overlay.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

# 4.6 Two-Stage
add_section(doc, "4.6 Model 4 — Two-Stage Screening Pipeline (EfficientNetV2B0 + Swin)", level=2)
add_section(doc, "4.6.1 Motivation", level=3)
add_para(doc,
    "The two-stage cascade mirrors the real-world clinical workflow: first "
    "detect whether DR is present (high-sensitivity screening), then grade "
    "severity only on positive cases. Stage 1 uses a lightweight CNN tuned for "
    "high recall; Stage 2 uses a powerful hierarchical transformer for "
    "fine-grained 4-class severity grading on the DR-only subset.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "4.6.2 Stage 1 — EfficientNetV2B0 Binary Screener", level=3)
add_bullets(doc, [
    "Framework: TensorFlow / Keras",
    "Base: EfficientNetV2B0 pretrained on ImageNet, include_top=False",
    "Head: GlobalAveragePooling2D → Dropout(0.3) → Dense(1, sigmoid)",
    "Label mapping: No_DR → 0, anything else → 1",
    "Augmentation: rescale 1/255, rotation 20°, horizontal flip, zoom 0.2",
    "Phase 1: head-only, Adam (default LR=1e-3), 5 epochs",
    "Phase 2: full fine-tuning, Adam LR=1e-5, 10 epochs",
    "Metrics: Accuracy, Recall (sensitivity), AUC",
    "Decision rule: if P(DR) < 0.3 → return No_DR (sensitivity-biased threshold)",
    "Saved artifact: efficientnetv2b0_binary_dr.h5",
])
add_section(doc, "4.6.3 Stage 2 — Swin Transformer 4-Class Severity Grader", level=3)
add_bullets(doc, [
    "Framework: PyTorch / timm",
    "Base: swin_base_patch4_window7_224 with replaced 4-class head",
    "Labels: Mild = 0, Moderate = 1, Severe = 2, Proliferate_DR = 3 (No_DR removed)",
    "Augmentation: Resize 224×224, RandomHorizontalFlip, RandomRotation ±10°, ImageNet normalization",
    "Phase 1: head-only, Adam LR=1e-3, 5 epochs",
    "Phase 2: full fine-tuning, Adam LR=1e-5, 15 epochs",
    "Loss: CrossEntropyLoss",
    "Saved artifacts: swin_checkpoint.pth, swin_transformer_dr.pth, complete_swin_model.pth",
])
add_section(doc, "4.6.4 Pipeline Fusion (Inference)", level=3)
add_para(doc,
    "predict_dr(image) → Stage 1 returns P(DR). If P(DR) < 0.3 return No_DR. "
    "Otherwise pass the image to Stage 2 which returns one of "
    "{Mild, Moderate, Severe, Proliferate_DR}. The low Stage-1 threshold "
    "biases the cascade toward sensitivity, in line with clinical screening priorities.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "4.7 System Workflow", level=2)
add_bullets(doc, [
    "4.7.1 Input Acquisition — clinician uploads a JPEG/PNG fundus image through the Flask interface.",
    "4.7.2 Preprocessing — image is resized to 224×224 and normalized using the model-specific preprocessing pipeline.",
    "4.7.3 DR Classification — the selected model (default: Vision Transformer, best validation accuracy) outputs class probabilities; argmax produces the predicted severity.",
    "4.7.4 Explainability — a GradCAM (TF models) or GradCAM++ (PyTorch transformers) heatmap is generated to highlight lesion-associated regions.",
    "4.7.5 Result Display — the result page shows the predicted severity badge, confidence percentage, GradCAM overlay, and clinical follow-up recommendation.",
])


# ===========================================================================
# CHAPTER 5 — RESULTS & DISCUSSION
# ===========================================================================
add_chapter(doc, 5, "RESULTS AND DISCUSSION")

add_section(doc, "5.1 Dataset and Experimental Setup", level=2)
add_para(doc,
    "All four models were trained on the same stratified 80/20 split of the "
    "APTOS 2019 dataset (3,662 images). Training: 2,929 images. Validation: "
    "733 images. Images were resized to 224×224 pixels. Models were trained "
    "on a single NVIDIA Tesla T4 GPU.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc, "Table 5.1: Training Parameters per Model", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
train_rows = [
    ("InceptionResNetV2", "2,929 (80 %)", "224 × 224", "20 (5 + 15)", "32"),
    ("Swin Transformer", "2,929 (80 %)", "224 × 224", "20 (5 + 15)", "32"),
    ("Vision Transformer", "2,929 (80 %)", "224 × 224", "15 (+2 phase 2)", "32"),
    ("Two-Stage (EffNetV2B0 + Swin)", "2,929 (80 %)", "224 × 224", "15 (S1) + 20 (S2)", "32"),
]
add_table(doc, ["Model", "Train Images", "Image Resolution", "Epochs", "Batch Size"], train_rows)

add_para(doc, "Table 5.2: Testing Parameters per Model", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
test_rows = [
    ("InceptionResNetV2", "731 (20 %)", "224 × 224", "32", "79.75 %"),
    ("Swin Transformer", "733 (20 %)", "224 × 224", "32", "84.72 %"),
    ("Vision Transformer", "733 (20 %)", "224 × 224", "32", "85.40 %"),
    ("Two-Stage (Stage 2 only)", "372 (DR-only)", "224 × 224", "32", "70.00 %"),
]
add_table(doc, ["Model", "Val Images", "Image Resolution", "Batch Size", "Best Val Accuracy"], test_rows)

add_section(doc, "5.2 Evaluation Metrics", level=2)
add_bullets(doc, [
    "Accuracy = (TP + TN) / Total — fraction of correctly classified images.",
    "Precision = TP / (TP + FP) — measures correctness of positive predictions.",
    "Recall (Sensitivity) = TP / (TP + FN) — measures the fraction of true positives detected.",
    "F1-Score = 2 × (Precision × Recall) / (Precision + Recall) — harmonic mean of precision and recall.",
    "ROC-AUC — computed per class in a One-vs-Rest setting using softmax probabilities.",
    "Confusion Matrix — 5×5 (or 4×4 for Stage 2) heatmap showing true vs. predicted class counts.",
])

# ---- 5.3 InceptionResNetV2 results
add_section(doc, "5.3 InceptionResNetV2 — Results", level=2)
add_para(doc,
    "The InceptionResNetV2 model was trained over a total of 20 epochs (5-epoch "
    "head warm-up + 15-epoch full fine-tuning). EarlyStopping was active "
    "(patience=5, monitor=val_loss). The best validation accuracy reached "
    "79.75 % on the 731-image validation subset.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc, "Table 5.3: InceptionResNetV2 — Per-Class Metrics", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
irv2_rows = [
    ("No DR", "0.94", "0.99", "0.97", "361"),
    ("Mild NPDR", "0.62", "0.61", "0.61", "74"),
    ("Moderate NPDR", "0.65", "0.80", "0.72", "199"),
    ("Severe NPDR", "0.35", "0.16", "0.22", "38"),
    ("Proliferative DR", "0.56", "0.17", "0.26", "59"),
]
add_table(doc, ["Class", "Precision", "Recall", "F1-Score", "Support"], irv2_rows)
add_para(doc, "Table 5.4: InceptionResNetV2 — Overall Performance", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
irv2_ov = [
    ("Accuracy", "0.7975"),
    ("Macro Precision", "0.62"),
    ("Macro Recall", "0.55"),
    ("Macro F1-Score", "0.56"),
    ("Weighted Precision", "0.77"),
    ("Weighted Recall", "0.79"),
    ("Weighted F1-Score", "0.77"),
]
add_table(doc, ["Metric", "Value"], irv2_ov, col_widths=[Cm(7), Cm(4)])
add_para(doc, "Figures: 5.1 Training Accuracy / Loss Curves (Fig 8), 5.2 Confusion Matrix (Fig 9), 5.3 ROC-AUC (Fig 10), 5.4 GradCAM Heatmap (Fig 11).", italic=True)
add_para(doc,
    "Discussion: The model excels at the dominant No_DR class (F1 = 0.97) and "
    "performs moderately on Moderate NPDR (F1 = 0.72) but struggles on minority "
    "classes — Severe NPDR (F1 = 0.22) and Proliferative DR (F1 = 0.26) — "
    "primarily due to the heavy class imbalance and the absence of class "
    "weighting or oversampling. Off-diagonal entries in the confusion matrix "
    "are concentrated between adjacent severity grades, which is clinically "
    "expected.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

# ---- 5.4 Swin
add_section(doc, "5.4 Swin Transformer — Results", level=2)
add_para(doc,
    "The Swin Transformer reached a peak validation accuracy of 84.72 % during "
    "Stage-2 fine-tuning, with the final-epoch validation accuracy of 83.49 % "
    "on 733 validation images.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc, "Table 5.5: Swin Transformer — Per-Class Metrics", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
swin_rows = [
    ("No DR", "0.97", "0.99", "0.98", "344"),
    ("Mild NPDR", "0.65", "0.76", "0.70", "79"),
    ("Moderate NPDR", "0.74", "0.79", "0.77", "193"),
    ("Severe NPDR", "0.72", "0.57", "0.64", "68"),
    ("Proliferative DR", "0.61", "0.39", "0.47", "49"),
]
add_table(doc, ["Class", "Precision", "Recall", "F1-Score", "Support"], swin_rows)
add_para(doc, "Table 5.6: Swin Transformer — Overall Performance", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
swin_ov = [
    ("Accuracy", "0.8347"),
    ("Macro Precision", "0.74"),
    ("Macro Recall", "0.70"),
    ("Macro F1-Score", "0.71"),
    ("Weighted Precision", "0.83"),
    ("Weighted Recall", "0.83"),
    ("Weighted F1-Score", "0.83"),
]
add_table(doc, ["Metric", "Value"], swin_ov, col_widths=[Cm(7), Cm(4)])
add_para(doc, "Figures: 5.5 Training Loss & Val Accuracy (Fig 12), 5.6 Confusion Matrix (Fig 13), 5.7 ROC-AUC (Fig 14), 5.8 GradCAM Samples (Fig 15).", italic=True)
add_para(doc,
    "Discussion: Swin's hierarchical shifted-window attention provides clear "
    "improvements over the CNN baseline on minority classes — Severe (F1 = 0.64) "
    "and Proliferative DR (F1 = 0.47) — while maintaining near-perfect No_DR "
    "performance. GradCAM heatmaps align well with clinically relevant lesion "
    "regions including microaneurysms, hemorrhages, and neovascularization.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

# ---- 5.5 ViT
add_section(doc, "5.5 Vision Transformer — Results", level=2)
add_para(doc,
    "The Vision Transformer achieved the highest overall validation accuracy "
    "of 85.40 % (626 / 733 correctly classified) on the held-out validation set, "
    "establishing it as the best-performing standalone model in this study.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc, "Table 5.7: Vision Transformer — Per-Class Metrics", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
vit_rows = [
    ("No DR", "0.98", "0.98", "0.98", "361"),
    ("Mild NPDR", "0.68", "0.46", "0.55", "74"),
    ("Moderate NPDR", "0.74", "0.89", "0.81", "200"),
    ("Severe NPDR", "0.61", "0.51", "0.56", "39"),
    ("Proliferative DR", "0.81", "0.66", "0.73", "59"),
]
add_table(doc, ["Class", "Precision", "Recall", "F1-Score", "Support"], vit_rows)
add_para(doc, "Table 5.8: Vision Transformer — Overall Performance", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
vit_ov = [
    ("Accuracy", "0.8540"),
    ("Macro Precision", "0.76"),
    ("Macro Recall", "0.70"),
    ("Macro F1-Score", "0.72"),
    ("Weighted Precision", "0.85"),
    ("Weighted Recall", "0.85"),
    ("Weighted F1-Score", "0.85"),
]
add_table(doc, ["Metric", "Value"], vit_ov, col_widths=[Cm(7), Cm(4)])
add_para(doc, "Figures: 5.9 Training Curves (Fig 16), 5.10 Confusion Matrix (Fig 17), 5.11 ROC-AUC (Fig 18), 5.12 GradCAM++ Samples (Fig 19).", italic=True)
add_para(doc,
    "Discussion: Global self-attention from layer 1, combined with "
    "WeightedRandomSampler and label smoothing, produces the most balanced "
    "per-class performance of all four models. Proliferative DR detection "
    "improves substantially (F1 = 0.73) compared to InceptionResNetV2 "
    "(0.26) and Swin (0.47), confirming the benefit of long-range attention "
    "for spatially distributed pathologies. A modest train/val loss gap "
    "(0.45 vs. 0.79) indicates mild overfitting controlled by label smoothing "
    "and gradient clipping.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

# ---- 5.6 Two-stage
add_section(doc, "5.6 Two-Stage Screening Pipeline — Results", level=2)
add_para(doc,
    "The Stage-2 Swin classifier was evaluated on the 372 DR-only validation "
    "images (No_DR samples removed). Stage-1 binary screening is evaluated "
    "via accuracy, recall (sensitivity), and AUC during training.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc, "Table 5.9: Two-Stage Pipeline — Stage 2 Per-Class Metrics", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
ts_rows = [
    ("Mild NPDR", "0.68", "0.69", "0.68", "74"),
    ("Moderate NPDR", "0.73", "0.84", "0.78", "200"),
    ("Severe NPDR", "0.54", "0.36", "0.43", "39"),
    ("Proliferative DR", "0.67", "0.47", "0.55", "59"),
]
add_table(doc, ["Class", "Precision", "Recall", "F1-Score", "Support"], ts_rows)
add_para(doc, "Table 5.10: Two-Stage Pipeline — Overall (Stage 2) Performance", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
ts_ov = [
    ("Accuracy (4-class)", "0.70"),
    ("Macro Precision", "0.65"),
    ("Macro Recall", "0.59"),
    ("Macro F1-Score", "0.61"),
    ("Weighted Precision", "0.69"),
    ("Weighted Recall", "0.70"),
    ("Weighted F1-Score", "0.69"),
]
add_table(doc, ["Metric", "Value"], ts_ov, col_widths=[Cm(7), Cm(4)])
add_para(doc, "Figures: 5.13 Stage-2 Confusion Matrix (Fig 21), 5.14 ROC-AUC (Fig 22).", italic=True)
add_para(doc,
    "Discussion: Removing the dominant No_DR class makes the Stage-2 problem "
    "harder than the equivalent 5-class problem because the model can no "
    "longer rely on the easy majority class to inflate weighted scores. The "
    "low Stage-1 threshold (0.3) biases the cascade toward sensitivity, "
    "matching clinical screening priorities. Severe NPDR remains the most "
    "challenging class due to its low sample count.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

# ---- 5.7 Comparative
add_section(doc, "5.7 Comparative Analysis", level=2)
add_para(doc, "Table 5.11: Comparative Performance of All Models (on 733 validation images, 5-class)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
cmp_rows = [
    ("InceptionResNetV2", "79.75 %", "0.62", "0.55", "0.56", "0.77"),
    ("Swin Transformer", "83.47 %", "0.74", "0.70", "0.71", "0.83"),
    ("Vision Transformer", "85.40 %", "0.76", "0.70", "0.72", "0.85"),
    ("Two-Stage (Stage 2, 4-class)", "70.00 %", "0.65", "0.59", "0.61", "0.69"),
]
add_table(doc, ["Model", "Val Acc.", "Macro P", "Macro R", "Macro F1", "Weighted F1"], cmp_rows)
add_para(doc,
    "Figure 5.15: Comparative Bar Chart of Model Accuracies (Fig 23).",
    italic=True,
)
add_para(doc,
    "Key takeaways from the comparative analysis:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_bullets(doc, [
    "Vision Transformer is the top-performing standalone model with 85.40 % validation accuracy and the best macro-F1 (0.72).",
    "Swin Transformer is a close second (83.47 %) with strong performance on minority classes.",
    "InceptionResNetV2 produces near-perfect No_DR classification but lags significantly on minority severity classes.",
    "The Two-Stage cascade is designed for clinical sensitivity rather than peak per-class accuracy; its 4-class Stage-2 accuracy is therefore not directly comparable to the 5-class scores of the other models.",
    "Transformer architectures (Swin, ViT) consistently outperform the CNN baseline on the 5-class severity task, particularly on minority classes (Severe, Proliferative DR).",
])


# ===========================================================================
# CHAPTER 6 — CONCLUSION
# ===========================================================================
add_chapter(doc, 6, "CONCLUSION AND FUTURE SCOPE")
add_section(doc, "6.1 Conclusion", level=2)
add_para(doc,
    "This project demonstrates an innovative and clinically motivated approach "
    "to addressing the global challenge of preventable vision loss from "
    "Diabetic Retinopathy using deep learning. Four architectures were "
    "implemented end-to-end on the APTOS 2019 dataset and evaluated using a "
    "consistent stratified 80/20 split: InceptionResNetV2 with multistage "
    "transfer learning (79.75 %), Swin Transformer with hierarchical shifted-"
    "window self-attention (83.47 %), Vision Transformer with single-phase "
    "fine-tuning, weighted sampling, and label smoothing (85.40 %), and a "
    "two-stage EfficientNetV2B0 + Swin cascade.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc,
    "The Vision Transformer emerged as the best-performing standalone model "
    "with the highest accuracy and the best macro-F1 score, particularly on "
    "minority severity classes. GradCAM / GradCAM++ explainability maps were "
    "produced for all transformer models, providing lesion-level "
    "interpretability that aligns with clinically relevant retinal markers — "
    "microaneurysms, hemorrhages, exudates, and neovascularization. Finally, a "
    "Flask web interface was developed to expose the trained model for "
    "real-time inference and visualization, making the system practical for "
    "deployment in low-resource clinical settings.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_para(doc,
    "In conclusion, this study confirms that transformer-based architectures, "
    "particularly the Vision Transformer with carefully tuned regularization "
    "and class-balancing strategies, achieve clinically competitive performance "
    "on multi-class DR severity grading while remaining interpretable through "
    "GradCAM-style attention visualization.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

add_section(doc, "6.2 Future Scope", level=2)
add_bullets(doc, [
    "Train an explicit ensemble of the three standalone models (soft / weighted / stacked voting) to combine their complementary strengths.",
    "Apply retinal-specific preprocessing (CLAHE, green-channel enhancement, Gaussian filtering) to all four models for a fair head-to-head comparison.",
    "Use larger and more diverse datasets (EyePACS, MESSIDOR-2) to improve generalization across fundus camera types and patient demographics.",
    "Develop a mobile application (React Native / Flutter) for point-of-care DR screening in rural clinics without internet connectivity.",
    "Integrate with Electronic Health Record (EHR) systems for automated screening alerts and longitudinal patient monitoring.",
    "Add an independent test set and external validation cohort for unbiased performance reporting.",
    "Replace the fixed Stage-1 threshold (0.3) in the two-stage pipeline with a learned operating point based on ROC analysis.",
    "Explore higher input resolutions (384×384 or 512×512) for transformer models to preserve fine-grained microaneurysm detail.",
])


# ===========================================================================
# CHAPTER 7 — PAPERS PUBLISHED
# ===========================================================================
add_chapter(doc, 7, "PAPERS PUBLISHED")
add_section(doc, "7.1 Papers Published", level=2)
add_para(doc,
    "A research paper based on the comparative evaluation of CNN and "
    "transformer architectures for diabetic retinopathy detection has been "
    "prepared and submitted for review. Citation details will be updated upon "
    "acceptance.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
)
add_section(doc, "7.2 Progress Reports Submitted to the Department", level=2)
add_bullets(doc, [
    "Progress Report I — Project synopsis, problem statement, and literature survey (submitted at the end of Phase-I).",
    "Progress Report II — Implementation milestones for InceptionResNetV2 and Swin Transformer.",
    "Progress Report III — Vision Transformer and Two-Stage pipeline implementation with preliminary metrics.",
    "Final Progress Report — Comparative evaluation, Flask integration, and final report submission (this document).",
])


# ===========================================================================
# CHAPTER 8 — REFERENCES
# ===========================================================================
add_chapter(doc, 8, "REFERENCES")
refs = [
    "[1] S. Guefrachi, et al., \"Diabetic Retinopathy Detection Using Deep Learning Multistage Training Method,\" Arabian Journal for Science and Engineering, Springer, 2025. DOI: 10.1007/s13369-024-09137-9",
    "[2] K. Balasamy and S. Suganyadevi, \"Multi-dimensional Fuzzy Based Diabetic Retinopathy Detection in Retinal Images through Deep CNN Method,\" Multimedia Tools and Applications, vol. 84, pp. 19625-19645, Springer, 2025. DOI: 10.1007/s11042-024-19798-1",
    "[3] R. Chandra, S. Tiwari, S. S. Kumar, and S. Agarwal, \"Diabetic Retinopathy Prediction Based on CNN and AlexNet Model,\" Indian Institute of Information Technology Allahabad, Prayagraj, India.",
    "[4] V. Gulshan, L. Peng, M. Coram, et al., \"Development and Validation of a Deep Learning Algorithm for Detection of Diabetic Retinopathy in Retinal Fundus Photographs,\" JAMA, vol. 316, no. 22, pp. 2402-2410, 2016.",
    "[5] M. D. Abramoff, et al., \"Pivotal Trial of an Autonomous AI-Based Diagnostic System for Detection of Diabetic Retinopathy in Primary Care Offices,\" npj Digital Medicine, vol. 1, no. 39, 2018.",
    "[6] H. Pratt, F. Coenen, D. M. Broadbent, S. P. Harding, and Y. Zheng, \"Convolutional Neural Networks for Diabetic Retinopathy,\" Procedia Computer Science, vol. 90, pp. 200-205, 2016.",
    "[7] M. Chetoui and M. A. Akhloufi, \"Explainable End-to-End Deep Learning for Diabetic Retinopathy Detection,\" Journal of Medical Imaging, vol. 7, no. 4, 2020.",
    "[8] A. Solanki, U. Singh, and S. Solanki, \"A Systematic Review on Diabetic Retinopathy Detection Using Deep Learning Techniques,\" Archives of Computational Methods in Engineering, 2023.",
    "[9] T. Li, et al., \"Diagnostic Assessment of Deep Learning Algorithms for Diabetic Retinopathy Screening,\" Information Sciences, vol. 501, pp. 511-522, 2019.",
    "[10] S. Wan, Y. Liang, and Y. Zhang, \"Deep Convolutional Neural Networks for Diabetic Retinopathy Detection by Image Classification,\" Computers & Electrical Engineering, vol. 72, pp. 274-282, 2018.",
    "[11] A. Dosovitskiy, L. Beyer, A. Kolesnikov, et al., \"An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale,\" International Conference on Learning Representations (ICLR), 2021.",
    "[12] Z. Liu, Y. Lin, Y. Cao, H. Hu, et al., \"Swin Transformer: Hierarchical Vision Transformer using Shifted Windows,\" IEEE/CVF International Conference on Computer Vision (ICCV), 2021.",
    "[13] M. Tan and Q. V. Le, \"EfficientNetV2: Smaller Models and Faster Training,\" Proceedings of the 38th International Conference on Machine Learning (ICML), 2021.",
    "[14] C. Szegedy, S. Ioffe, V. Vanhoucke, and A. Alemi, \"Inception-v4, Inception-ResNet and the Impact of Residual Connections on Learning,\" AAAI Conference on Artificial Intelligence, 2017.",
    "[15] R. R. Selvaraju, M. Cogswell, A. Das, R. Vedantam, D. Parikh, and D. Batra, \"Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization,\" International Journal of Computer Vision, vol. 128, pp. 336-359, 2020.",
    "[16] A. Chattopadhay, A. Sarkar, P. Howlader, and V. N. Balasubramanian, \"Grad-CAM++: Improved Visual Explanations for Deep Convolutional Networks,\" IEEE Winter Conference on Applications of Computer Vision (WACV), 2018.",
    "[17] APTOS 2019 Blindness Detection Dataset, Kaggle, https://www.kaggle.com/c/aptos2019-blindness-detection/data",
    "[18] R. Wightman, \"PyTorch Image Models (timm),\" https://github.com/rwightman/pytorch-image-models, 2019-present.",
    "[19] M. Abadi et al., \"TensorFlow: Large-Scale Machine Learning on Heterogeneous Distributed Systems,\" Google Research White Paper, 2016.",
    "[20] A. Paszke et al., \"PyTorch: An Imperative Style, High-Performance Deep Learning Library,\" Advances in Neural Information Processing Systems (NeurIPS), 2019.",
]
for ref in refs:
    add_para(doc, ref, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
